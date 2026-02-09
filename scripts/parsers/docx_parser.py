"""
DOCX Parser

使用 python-docx 将 DOCX 简历转换为 Markdown 格式文本。

支持：
- 标题层级 (Heading 1-6 -> # ~ ######)
- 段落文本 (粗体/斜体 -> **bold** / *italic*)
- 列表 (有序/无序)
- 表格 (Markdown table 格式)
"""

import re
from pathlib import Path
from typing import List, Optional

try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    Document = None
    WD_ALIGN_PARAGRAPH = None


class DocxParser:
    """DOCX 文档解析器"""

    def parse(self, file_path: str) -> str:
        """解析 DOCX 文件，返回 Markdown 格式文本。

        Args:
            file_path: DOCX 文件路径。

        Returns:
            Markdown 格式的纯文本内容。

        Raises:
            ImportError: 未安装 python-docx。
            FileNotFoundError: 文件不存在。
            ValueError: 文件不是有效的 DOCX。
        """
        if Document is None:
            raise ImportError(
                "python-docx is required for DOCX parsing. "
                "Install it with: pip install python-docx"
            )

        path = Path(file_path)
        if not path.exists():
            raise FileNotFoundError(f"DOCX file not found: {file_path}")
        if path.suffix.lower() not in (".docx", ".doc"):
            raise ValueError(f"Not a DOCX file: {file_path}")

        doc = Document(str(path))
        parts: List[str] = []

        # 遍历文档的 body 元素，保持原始顺序
        # python-docx 的 document.element.body 包含段落和表格的混合序列
        for element in doc.element.body:
            tag = element.tag.split("}")[-1]  # 去掉命名空间

            if tag == "p":
                # 段落元素 - 查找对应的 Paragraph 对象
                para = self._find_paragraph(doc, element)
                if para is not None:
                    md = self._convert_paragraph(para)
                    if md is not None:
                        parts.append(md)

            elif tag == "tbl":
                # 表格元素 - 查找对应的 Table 对象
                table = self._find_table(doc, element)
                if table is not None:
                    md = self._convert_table(table)
                    if md:
                        parts.append(md)

        return self._clean_output("\n\n".join(parts))

    def _find_paragraph(self, doc, element):
        """根据 XML 元素查找对应的 Paragraph 对象。"""
        for para in doc.paragraphs:
            if para._element is element:
                return para
        return None

    def _find_table(self, doc, element):
        """根据 XML 元素查找对应的 Table 对象。"""
        for table in doc.tables:
            if table._element is element:
                return table
        return None

    def _convert_paragraph(self, para) -> Optional[str]:
        """将段落转换为 Markdown 文本。

        处理：标题、列表项、普通段落（含粗体/斜体）。
        """
        style_name = (para.style.name or "").lower()
        text = self._runs_to_markdown(para.runs)

        if not text.strip():
            return None

        # 标题
        if style_name.startswith("heading"):
            level = self._heading_level(style_name)
            prefix = "#" * level
            return f"{prefix} {text.strip()}"

        # 列表 - 通过样式名检测
        if "list bullet" in style_name or "list" in style_name and "bullet" in style_name:
            return f"- {text.strip()}"
        if "list number" in style_name:
            return f"1. {text.strip()}"

        # 检查段落是否有缩进 + 项目符号 XML 标记
        if self._is_bulleted(para):
            return f"- {text.strip()}"
        if self._is_numbered(para):
            return f"1. {text.strip()}"

        return text

    def _runs_to_markdown(self, runs) -> str:
        """将段落中的 runs 转换为带格式的 Markdown 文本。

        Run 是具有统一格式的文本片段。
        """
        parts: List[str] = []
        for run in runs:
            text = run.text
            if not text:
                continue

            is_bold = run.bold
            is_italic = run.italic

            if is_bold and is_italic:
                parts.append(f"***{text}***")
            elif is_bold:
                parts.append(f"**{text}**")
            elif is_italic:
                parts.append(f"*{text}*")
            else:
                parts.append(text)

        return "".join(parts)

    def _heading_level(self, style_name: str) -> int:
        """从样式名中提取标题级别。

        'heading 1' -> 1, 'heading 2' -> 2, etc.
        """
        match = re.search(r"(\d+)", style_name)
        if match:
            level = int(match.group(1))
            return min(max(level, 1), 6)
        return 1

    def _is_bulleted(self, para) -> bool:
        """检查段落是否为无序列表项（通过 XML 标记检测）。"""
        pPr = para._element.find(
            "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pPr"
        )
        if pPr is None:
            return False
        numPr = pPr.find(
            "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}numPr"
        )
        if numPr is None:
            return False
        # 有 numPr 但没有 ilvl > 0 且有 numId，可能是列表
        numId = numPr.find(
            "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}numId"
        )
        if numId is not None:
            val = numId.get(
                "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val"
            )
            # numId=0 通常表示不是列表
            if val and val != "0":
                return True
        return False

    def _is_numbered(self, para) -> bool:
        """检查段落是否为有序列表项。

        注意：python-docx 中无序和有序列表都通过 numPr 标记，
        精确区分需要查询 numbering.xml。此处简化处理：
        _is_bulleted 已覆盖所有带 numPr 的情况，
        这个方法作为备用检测。
        """
        return False

    def _convert_table(self, table) -> str:
        """将表格转换为 Markdown table 格式。"""
        rows = table.rows
        if not rows:
            return ""

        md_rows: List[str] = []

        for i, row in enumerate(rows):
            cells = [self._cell_text(cell) for cell in row.cells]
            md_rows.append("| " + " | ".join(cells) + " |")

            # 第一行后添加分隔线
            if i == 0:
                separator = "| " + " | ".join("---" for _ in cells) + " |"
                md_rows.append(separator)

        return "\n".join(md_rows)

    def _cell_text(self, cell) -> str:
        """提取单元格文本，合并多段落为单行。"""
        texts = []
        for para in cell.paragraphs:
            text = para.text.strip()
            if text:
                texts.append(text)
        # 单元格内多段落用 <br> 连接
        return " <br> ".join(texts) if texts else ""

    def _clean_output(self, text: str) -> str:
        """清理最终输出。

        - 合并连续空行
        - 修剪首尾
        - 合并相邻的格式标记 (如 **text****more** -> **textmore**)
        """
        # 合并相邻的粗体标记
        text = re.sub(r"\*\*\*\*", "", text)
        # 合并相邻的斜体标记
        text = re.sub(r"\*\*(?!\*)", "**", text)
        # 合并连续空行
        text = re.sub(r"\n{3,}", "\n\n", text)
        return text.strip()
